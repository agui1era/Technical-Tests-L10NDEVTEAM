import argparse
import json
import os
from docx import Document

# ---------- Extractor: DOCX → JSON ----------
def extract_docx_to_json(input_docx, output_json):
    data = []
    doc = Document(input_docx)

    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            data.append({
                "id": f"{os.path.basename(input_docx)}_{i}",
                "key": f"seg_{i}",
                "source": para.text.strip(),
                "target": ""
            })

    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    print(f"✅ Extracted {len(data)} segments to {output_json}")

# ---------- Builder: JSON → DOCX ----------
def build_json_to_docx(input_json, output_docx):
    with open(input_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    for seg in data:
        text = seg.get("target") or seg.get("source")
        doc.add_paragraph(text)

    doc.save(output_docx)
    print(f"✅ Built DOCX at {output_docx}")

# ---------- Main CLI ----------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="DOCX ⇆ JSON Parser")
    parser.add_argument("mode", choices=["extract", "build"], help="Operation mode")
    parser.add_argument("input", help="Input file (.docx or .json)")
    parser.add_argument("output", help="Output file (.json or .docx)")
    args = parser.parse_args()

    if args.mode == "extract":
        extract_docx_to_json(args.input, args.output)
    elif args.mode == "build":
        build_json_to_docx(args.input, args.output)